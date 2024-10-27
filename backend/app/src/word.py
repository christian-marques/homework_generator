import os
from docx import Document
from lxml import etree
from datetime import datetime
from .db import students_school_year

# Função para obter a data atual no formato aaaammdd
def get_current_date(str_format):
    return datetime.now().strftime(str_format)

# Função para gerar o nome do arquivo baseado no conteúdo
def generate_filename(student_name, class_name, theme):
    current_date = get_current_date("%Y-%m-%d")
    # Gera o nome do arquivo no formato aaaammdd_NomeAluno_Disciplina_Tema.docx
    filename = f"{current_date}_{student_name.replace(' ', '_')}_{class_name.replace(' ', '_')}_{theme.replace(' ', '_')}.docx"
    return filename

# Função para substituir o placeholder em um parágrafo ou célula
def replace_placeholders_in_paragraph(paragraph, placeholder, replacement):
    if placeholder in paragraph.text:
        paragraph.text = paragraph.text.replace(placeholder, replacement)

# Função para substituir os placeholders no documento, incluindo cabeçalho
def replace_placeholders_in_document(doc, placeholder, replacement):
    # Substitui nos parágrafos do corpo principal
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, placeholder, replacement)

    # Substitui em células de tabelas (caso tenha)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders_in_paragraph(cell.paragraphs[0], placeholder, replacement)

    # Substitui nos cabeçalhos
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            replace_placeholders_in_paragraph(paragraph, placeholder, replacement)
        # Acessar e substituir texto nas caixas de texto (shapes) no cabeçalho
        for shape in header._element.xpath('.//w:t'):
            if placeholder in shape.text:
                shape.text = shape.text.replace(placeholder, replacement)

# Função para gerar o arquivo Word
def generate_file(template_path, output_path, header_info, message):
    # Carrega o template Word
    doc = Document(template_path)

    # Preenche o documento com os dados dinâmicos
    student_name, class_name, theme = header_info
    replace_placeholders_in_document(doc, '{name}', student_name)
    replace_placeholders_in_document(doc, '{year}', str(students_school_year.get(student_name)))
    replace_placeholders_in_document(doc, '{class}', class_name)
    replace_placeholders_in_document(doc, 'theme', theme.upper())
    replace_placeholders_in_document(doc, '{date}', str(get_current_date("%d/%m/%Y")))
    
    # Para o enunciado, caso tenha múltiplas linhas
    replace_placeholders_in_document(doc, '{enunciado}', message)

    # Gera o nome do arquivo com base no conteúdo
    filename = generate_filename(student_name, class_name, theme)
    output_filepath = os.path.join(output_path, filename)

    print(f">>> output_filepath: {output_filepath}")

    # Salva o arquivo Word gerado
    doc.save(output_filepath)
    return output_filepath, filename